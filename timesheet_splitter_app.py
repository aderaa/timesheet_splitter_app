import os
import io
import re
import shutil
import zipfile
import json
from datetime import datetime

import pandas as pd
from docx import Document

import streamlit as st

# Try to import reportlab for PDF generation
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


# ==== Paths & Config ====

APP_ROOT = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
APP_CONFIG_FILE = os.path.join(APP_ROOT, "app.config")

# Columns in "1.Full Timesheet Report.xlsx"
FULL_EMP_ID_COL = "EMP ID"
FULL_EMP_NAME_COL = "User Name"
FULL_DATE_COL = "Date"
FULL_HOURS_COL = "Regular Time (Hours)"
FULL_PROJECT_TYPE_COL = "Project Type"  # column used for ignore list (e.g. "On Bench")

# Columns in "3.Databse.xlsx" / database file
DB_EMP_ID_COL = "Emp ID"
DB_EMP_NAME_COL = "Employee Name"
DB_VENDOR_COL = "Organization"

SUMMARY_DOC_NAME = "Timesheet_Split_Summary.docx"

# Default vendor for P* employees not in DB
DEFAULT_VENDOR_NAME = "UnAssigned Vendor"

# Default ignore list for Project Type
DEFAULT_IGNORE_LIST = ["On Bench"]

DEFAULT_CONFIG = {
    "database_path": "Database.xlsx",          # relative to app root
    "ignore_project_types": DEFAULT_IGNORE_LIST,
    "logo_path": "malomatia-logo.png",         # relative to app root
    "department_name": "Digital Services",
    "user_name": "Malomatian",
    "default_output_mode": "folder",           # "folder" or "zip"
    "default_output_folder": "output",
    "language": "English",                     # persisted language preference
}


# ==== Simple i18n helper ====

TEXT = {
    "en": {
        "title": "Timesheet Splitter for Outsourced Staff",
        "subtitle": "Upload the full Clarity timesheet export and the vendor database. The app will create one Excel + PDF file per employee, grouped under folders per vendor, plus a DOCX summary report.",
        "settings": "Settings",
        "output_mode": "Output mode",
        "output_mode_folder": "Save files to a folder on this machine",
        "output_mode_zip": "Download everything as a ZIP file",
        "output_folder": "Target output folder (relative or absolute path)",
        "full_timesheet": "1) Full Timesheet Report (from Clarity)",
        "vendor_db": "2) Vendor / Employee Database",
        "folder_not_empty": "⚠️ The folder '{path}' is not empty. Its contents will be deleted before export.",
        "confirm_clear": "I understand, clear this folder before export",
        "start": "🚀 Start splitting",
        "fatal_error": "Unexpected error",
        "progress_init": "Preparing data...",
        "progress_emp": "Processing employee {idx}/{total}: {emp_id}",
        "done": "✅ Done! Timesheets have been split successfully.",
        "download_summary": "Download summary (DOCX)",
        "download_zip": "Download ZIP file",
        "metrics_title": "Run summary",
        "metric_total_emps": "Unique employees in source",
        "metric_exported_emps": "Employees exported",
        "metric_ignored_emps": "Ignored (not in DB & not P*)",
        "metric_failed_emps": "Failed splits",
        "metric_unassigned_emps": "Unassigned 'P*' employees",
        "metric_project_flagged_emps": "Employees with ignored Project Type",
        "vendor_summary_title": "Hours per vendor",
        "exported_table_title": "Exported employees",
        "ignored_table_title": "Ignored employees (not in DB & not P*)",
        "failed_table_title": "Failed splits",
        "unassigned_table_title": "Unassigned 'P*' employees (default vendor)",
        "project_ignored_table_title": "Employees with Project Types in ignore list (NOT exported)",
        "no_failed": "No failed splits 🎉",
        "no_ignored": "No ignored employees 🎉",
        "no_unassigned": "No unassigned 'P*' employees 🎉",
        "no_project_flagged": "No employees with ignored Project Types 🎉",
        "run_timestamp": "Run timestamp",
        "period": "Timesheet period",
        "no_dates": "Not available",
        "summary_doc_title": "Timesheet Split Summary",
        "doc_section_overview": "1. Overview",
        "doc_section_vendor_summary": "2. Hours per vendor",
        "doc_section_exported_emps": "3. Exported employees",
        "doc_section_ignored": "4. Ignored employees (not in vendor database and not starting with 'P')",
        "doc_section_failed": "5. Failed splits",
        "doc_section_unassigned": "6. 'P*' employees not in vendor DB (assigned to default vendor)",
        "doc_section_project_ignored": "7. Employees with Project Types in ignore list (NOT exported)",
        "doc_overview_bullet_1": "Run timestamp: {ts}",
        "doc_overview_bullet_2": "Timesheet period: {period}",
        "doc_overview_bullet_3": "Total unique employees in source: {total}",
        "doc_overview_bullet_4": "Employees exported: {exported}",
        "doc_overview_bullet_5": "Ignored employees (not in DB & not P*): {ignored}",
        "doc_overview_bullet_6": "Failed splits: {failed}",
        "doc_overview_bullet_7": "Unassigned 'P*' employees (default vendor): {unassigned}",
        "doc_overview_bullet_8": "Employees with project types in ignore list (NOT exported): {project_flagged}",
        "table_vendor": "Vendor",
        "table_vendor_hours": "Total hours",
        "table_vendor_emp_count": "Employees",
        "table_emp_id": "Emp ID",
        "table_emp_name": "Employee Name",
        "table_emp_hours": "Total hours",
        "table_reason": "Reason",
        "table_unassigned_reason": "Reason",
        "table_project_types": "Ignored Project Types",
        "reason_not_in_db": "Not in vendor database",
        "reason_unassigned": "Emp ID starts with 'P' and not found in vendor database (assigned to default vendor)",
        "reason_exception": "Exception while writing file",
        "ui_language": "Language / اللغة",
        "ui_language_en": "English",
        "ui_language_ar": "العربية",
        "ignore_list": "Ignored Project Types (from config)",
        "pdf_not_available": "PDF generation library (reportlab) is not installed. Only Excel files will be generated. Please run: pip install reportlab",
        "vendor_staff_header_emp_count": "Employees count",
        "vendor_staff_header_total_hours": "Total hours",
        "vendor_staff_header_avg_hours": "Average hours per employee",
        "sidebar_page_main": "Timesheet Splitter",
        "sidebar_page_settings": "App Settings",
        "config_title": "Application Configuration",
        "config_section_paths": "Paths & Files",
        "config_db_path": "Vendor / Employee DB file path (relative to app root or absolute)",
        "config_logo_path": "Logo image path (relative to app root or absolute)",
        "config_section_branding": "Branding",
        "config_department_name": "Department name",
        "config_user_name": "Default user name",
        "config_section_output": "Output Defaults",
        "config_default_output_mode": "Default output mode",
        "config_default_output_folder": "Default output folder",
        "config_section_ignore": "Ignored Project Types (Project Type column)",
        "config_save_button": "💾 Save configuration",
        "config_saved": "Configuration saved.",
        "config_logo_preview": "Logo preview",
        "config_db_resolved": "Resolved DB path",
        "db_loaded_from_config": "Loaded vendor DB from config path:",
        "db_config_missing": "Configured DB file not found at:",
        "db_config_error": "Error reading DB file from configured path:",
        "db_upload_optional": "Override Vendor / Employee Database (optional)",
        "sidebar_lang": "Language / اللغة",
    },
    "ar": {
        "title": "أداة تقسيم الجداول الزمنية للموظفين الخارجيين",
        "subtitle": "قم برفع ملف الجداول الزمنية الكامل من Clarity وملف قاعدة بيانات الموردين، وستقوم الأداة بإنشاء ملف Excel وملف PDF لكل موظف داخل مجلد باسم المورد مع تقرير ملخص بصيغة DOCX.",
        "settings": "الإعدادات",
        "output_mode": "طريقة الإخراج",
        "output_mode_folder": "حفظ الملفات في مجلد على هذا الجهاز",
        "output_mode_zip": "تحميل كل الملفات كملف ZIP واحد",
        "output_folder": "مسار مجلد الإخراج (نسبي أو مطلق)",
        "full_timesheet": "١) ملف الجداول الزمنية الكامل (من Clarity)",
        "vendor_db": "٢) ملف قاعدة بيانات الموردين / الموظفين",
        "folder_not_empty": "⚠️ المجلد '{path}' غير فارغ. سيتم حذف محتوياته قبل التصدير.",
        "confirm_clear": "أقرّ بذلك، قم بإفراغ هذا المجلد قبل التصدير",
        "start": "🚀 ابدأ عملية التقسيم",
        "fatal_error": "خطأ غير متوقع",
        "progress_init": "جاري تجهيز البيانات...",
        "progress_emp": "جاري معالجة الموظف {idx} من {total}: {emp_id}",
        "done": "✅ تم التنفيذ! تم تقسيم الجداول الزمنية بنجاح.",
        "download_summary": "تحميل تقرير الملخص (DOCX)",
        "download_zip": "تحميل ملف ZIP",
        "metrics_title": "ملخص العملية",
        "metric_total_emps": "عدد الموظفين في المصدر",
        "metric_exported_emps": "عدد الموظفين الذين تم تصديرهم",
        "metric_ignored_emps": "الموظفون المتجاهَلون (غير موجودين بقاعدة البيانات ولا يبدأ رقمهم بـ P)",
        "metric_failed_emps": "عدد المحاولات الفاشلة",
        "metric_unassigned_emps": "موظفو P غير المربوطين بمورد",
        "metric_project_flagged_emps": "موظفون لديهم نوع مشروع من قائمة التجاهل",
        "vendor_summary_title": "ساعات العمل لكل مورد",
        "exported_table_title": "الموظفون الذين تم تصديرهم",
        "ignored_table_title": "الموظفون المتجاهَلون (غير موجودين بقاعدة البيانات ولا يبدأ رقمهم بـ P)",
        "failed_table_title": "المحاولات الفاشلة",
        "unassigned_table_title": "الموظفون الذين يبدأ رقمهم بـ P وغير موجودين بقاعدة البيانات (المورد الافتراضي)",
        "project_ignored_table_title": "الموظفون الذين لديهم أنواع مشروع ضمن قائمة التجاهل (لا يتم تصديرهم)",
        "no_failed": "لا توجد محاولات فاشلة 🎉",
        "no_ignored": "لا يوجد موظفون متجاهَلون 🎉",
        "no_unassigned": "لا يوجد موظفو P غير مربوطين بمورد 🎉",
        "no_project_flagged": "لا يوجد موظفون لديهم أنواع مشروع من قائمة التجاهل 🎉",
        "run_timestamp": "تاريخ ووقت التشغيل",
        "period": "فترة الجداول الزمنية",
        "no_dates": "غير متوفر",
        "summary_doc_title": "تقرير ملخص تقسيم الجداول الزمنية",
        "doc_section_overview": "١. نظرة عامة",
        "doc_section_vendor_summary": "٢. ساعات العمل لكل مورد",
        "doc_section_exported_emps": "٣. الموظفون الذين تم تصديرهم",
        "doc_section_ignored": "٤. الموظفون المتجاهَلون (غير موجودين في قاعدة بيانات الموردين ولا يبدأ رقمهم بـ P)",
        "doc_section_failed": "٥. المحاولات الفاشلة",
        "doc_section_unassigned": "٦. الموظفون الذين يبدأ رقمهم بـ P وغير موجودين في قاعدة بيانات الموردين (المورد الافتراضي)",
        "doc_section_project_ignored": "٧. الموظفون الذين لديهم أنواع مشروع ضمن قائمة التجاهل (لا يتم تصديرهم)",
        "doc_overview_bullet_1": "تاريخ ووقت التشغيل: {ts}",
        "doc_overview_bullet_2": "فترة الجداول الزمنية: {period}",
        "doc_overview_bullet_3": "إجمالي عدد الموظفين في المصدر: {total}",
        "doc_overview_bullet_4": "عدد الموظفين الذين تم تصديرهم: {exported}",
        "doc_overview_bullet_5": "عدد الموظفين المتجاهَلين (غير موجودين في قاعدة البيانات ولا يبدأ رقمهم بـ P): {ignored}",
        "doc_overview_bullet_6": "عدد المحاولات الفاشلة: {failed}",
        "doc_overview_bullet_7": "عدد موظفي P غير المربوطين بمورد (المورد الافتراضي): {unassigned}",
        "doc_overview_bullet_8": "عدد الموظفين الذين لديهم نوع مشروع ضمن قائمة التجاهل (لا يتم تصديرهم): {project_flagged}",
        "table_vendor": "المورد",
        "table_vendor_hours": "إجمالي الساعات",
        "table_vendor_emp_count": "عدد الموظفين",
        "table_emp_id": "رقم الموظف",
        "table_emp_name": "اسم الموظف",
        "table_emp_hours": "إجمالي الساعات",
        "table_reason": "السبب",
        "table_unassigned_reason": "السبب",
        "table_project_types": "أنواع المشروع من قائمة التجاهل",
        "reason_not_in_db": "غير موجود في قاعدة بيانات الموردين",
        "reason_unassigned": "رقم الموظف يبدأ بـ P وغير موجود في قاعدة بيانات الموردين (تم إسناده للمورد الافتراضي)",
        "reason_exception": "خطأ أثناء حفظ الملف",
        "ui_language": "Language / اللغة",
        "ui_language_en": "English",
        "ui_language_ar": "العربية",
        "ignore_list": "أنواع المشروع المتجاهَلة (من إعدادات التطبيق)",
        "pdf_not_available": "مكتبة إنشاء ملفات PDF (reportlab) غير مثبتة، سيتم إنشاء ملفات Excel فقط. برجاء تنفيذ الأمر: pip install reportlab",
        "vendor_staff_header_emp_count": "عدد الموظفين",
        "vendor_staff_header_total_hours": "إجمالي الساعات",
        "vendor_staff_header_avg_hours": "متوسط الساعات لكل موظف",
        "sidebar_page_main": "أداة التقسيم",
        "sidebar_page_settings": "إعدادات التطبيق",
        "config_title": "إعدادات التطبيق",
        "config_section_paths": "المسارات والملفات",
        "config_db_path": "مسار ملف قاعدة بيانات الموردين / الموظفين (نسبي من مجلد التطبيق أو مطلق)",
        "config_logo_path": "مسار شعار الجهة (نسبي من مجلد التطبيق أو مطلق)",
        "config_section_branding": "الهوية (Branding)",
        "config_department_name": "اسم الإدارة / القسم",
        "config_user_name": "اسم المستخدم الافتراضي",
        "config_section_output": "إعدادات الإخراج الافتراضية",
        "config_default_output_mode": "طريقة الإخراج الافتراضية",
        "config_default_output_folder": "مجلد الإخراج الافتراضي",
        "config_section_ignore": "أنواع المشروع المتجاهَلة (عمود Project Type)",
        "config_save_button": "💾 حفظ الإعدادات",
        "config_saved": "تم حفظ الإعدادات.",
        "config_logo_preview": "معاينة الشعار",
        "config_db_resolved": "المسار الفعلي لملف قاعدة البيانات",
        "db_loaded_from_config": "تم تحميل ملف قاعدة بيانات الموردين من مسار الإعدادات:",
        "db_config_missing": "لم يتم العثور على ملف قاعدة بيانات الموردين في المسار التالي:",
        "db_config_error": "خطأ أثناء قراءة ملف قاعدة البيانات من المسار المحدد:",
        "db_upload_optional": "رفع ملف قاعدة بيانات بديل (اختياري)",
        "sidebar_lang": "Language / اللغة",
    },
}


def t(key: str, lang: str) -> str:
    """Translation helper."""
    return TEXT.get(lang, TEXT["en"]).get(key, TEXT["en"].get(key, key))


# ==== Styling ====
CUSTOM_CSS = """
<style>
.app-title {
    text-align: left;
    color: #2c3e50;
    font-size: 2.3rem;
    margin-bottom: 0.2rem;
}
.app-subtitle {
    text-align: left;
    color: #555;
    margin-bottom: 0.3rem;
}
.brand-line {
    color: #333;
    font-size: 0.95rem;
    margin-bottom: 0.3rem;
}
.welcome-line {
    color: #2c3e50;
    font-size: 1.1rem;
    margin-bottom: 1.2rem;
}
.summary-card {
    padding: 0.9rem 1.2rem;
    border-radius: 12px;
    background: linear-gradient(135deg, #1abc9c, #3498db);
    color: white;
    margin-bottom: 1.2rem;
}
</style>
"""


# ==== Config helpers ====

def load_app_config() -> dict:
    cfg = DEFAULT_CONFIG.copy()
    if os.path.exists(APP_CONFIG_FILE):
        try:
            with open(APP_CONFIG_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                cfg.update(data)
        except Exception:
            pass
    return cfg


def save_app_config(cfg: dict) -> None:
    try:
        with open(APP_CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
    except Exception:
        pass


# ==== Core helpers ====

def safe_name(name: str) -> str:
    """Make a string safe for use as a Windows file/folder name."""
    if pd.isna(name):
        name = "Unknown"
    name = str(name)
    return re.sub(r'[<>:"/\\\\|?*]', "_", name)


def resolve_path_from_config(path_value: str) -> str:
    """Resolve a path from config: absolute or relative to APP_ROOT."""
    if not path_value:
        return ""
    path_value = str(path_value)
    if os.path.isabs(path_value):
        return path_value
    return os.path.join(APP_ROOT, path_value)


def dataframe_to_pdf_bytes(df: pd.DataFrame, title: str = "") -> bytes:
    """Render a pandas DataFrame to a table PDF, word-wrapped and fitted to landscape A4 width."""
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab is not installed")

    buffer = io.BytesIO()

    page_size = landscape(A4)
    page_width, page_height = page_size
    left_margin = right_margin = top_margin = bottom_margin = 20

    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
    )

    elements = []
    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        "HeaderCell",
        parent=styles["Normal"],
        fontSize=8,
        leading=10,
        alignment=1,  # center
    )
    cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["Normal"],
        fontSize=7,
        leading=9,
        wordWrap="CJK",
    )

    if title:
        elements.append(Paragraph(title, styles["Heading2"]))

    if df is None or df.empty:
        df = pd.DataFrame({"": ["(no data)"]})

    df_str = df.astype(str)

    header_row = [Paragraph(str(col), header_style) for col in df_str.columns]
    data_rows = []
    for _, row in df_str.iterrows():
        data_rows.append([Paragraph(str(val), cell_style) for val in row])

    data = [header_row] + data_rows

    ncols = len(df_str.columns) or 1
    avail_width = page_width - left_margin - right_margin
    col_widths = [avail_width / ncols] * ncols

    table = Table(data, colWidths=col_widths, repeatRows=1)

    style = TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("ALIGN", (0, 1), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
    )
    table.setStyle(style)

    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def prepare_employee_data(full_df: pd.DataFrame, db_df: pd.DataFrame, ignore_list):
    """
    Build in-memory structures for employees:
    - employees: list of dicts with vendor, emp_id, emp_name, hours, df, flags (only those to be exported)
    - ignored: employees present in full_df but not in db_df and not starting with 'P'
    - unassigned: P* employees not in db but assigned to default vendor
    - project_flagged: employees whose rows contain project types from ignore_list (NOT exported)
    """

    obj_cols = full_df.select_dtypes(include=["object"]).columns
    full_df[obj_cols] = full_df[obj_cols].ffill()

    if FULL_DATE_COL in full_df.columns:
        full_df[FULL_DATE_COL] = pd.to_datetime(full_df[FULL_DATE_COL], errors="coerce")
    if FULL_HOURS_COL in full_df.columns:
        full_df[FULL_HOURS_COL] = pd.to_numeric(full_df[FULL_HOURS_COL], errors="coerce").fillna(0.0)
    else:
        full_df[FULL_HOURS_COL] = 0.0

    db_df = db_df.copy()
    emp_id_to_vendor = {}
    emp_id_to_name = {}

    if DB_EMP_ID_COL not in db_df.columns or DB_VENDOR_COL not in db_df.columns:
        missing = []
        if DB_EMP_ID_COL not in db_df.columns:
            missing.append(DB_EMP_ID_COL)
        if DB_VENDOR_COL not in db_df.columns:
            missing.append(DB_VENDOR_COL)
        raise ValueError(
            f"Vendor DB is missing required columns: {', '.join(missing)}. "
            f"Found columns: {list(db_df.columns)}"
        )

    for _, row in db_df.iterrows():
        emp_id = str(row[DB_EMP_ID_COL]).strip()
        vendor = row[DB_VENDOR_COL]
        emp_name = row.get(DB_EMP_NAME_COL, None)
        if emp_id:
            emp_id_to_vendor[emp_id] = vendor
            if emp_name is not None and not pd.isna(emp_name):
                emp_id_to_name[emp_id] = emp_name

    ignore_set = {v.strip().lower() for v in ignore_list if v and isinstance(v, str)}

    employees = []
    ignored = []
    unassigned = []
    project_flagged = []

    if FULL_EMP_ID_COL not in full_df.columns or FULL_EMP_NAME_COL not in full_df.columns:
        raise ValueError("Full timesheet file is missing required columns.")

    unique_emp_ids = sorted({str(x).strip() for x in full_df[FULL_EMP_ID_COL].dropna().unique()})
    has_project_type_col = FULL_PROJECT_TYPE_COL in full_df.columns

    for emp_id in unique_emp_ids:
        emp_rows = full_df[full_df[FULL_EMP_ID_COL].astype(str).str.strip() == emp_id]
        if emp_rows.empty:
            continue

        emp_name_series = emp_rows[FULL_EMP_NAME_COL].dropna()
        emp_name = str(emp_name_series.iloc[0]) if not emp_name_series.empty else ""
        total_hours = float(emp_rows[FULL_HOURS_COL].sum())

        is_unassigned = False

        if emp_id not in emp_id_to_vendor:
            if emp_id.startswith("P"):
                vendor = DEFAULT_VENDOR_NAME
                is_unassigned = True
                unassigned.append(
                    {
                        "Vendor": vendor,
                        "Emp ID": emp_id,
                        "Employee Name": emp_name,
                        "Total Hours": total_hours,
                    }
                )
            else:
                ignored.append(
                    {
                        "Emp ID": emp_id,
                        "Employee Name": emp_name,
                        "Total Hours": total_hours,
                        "Reason": "not_in_db",
                    }
                )
                continue
        else:
            vendor = emp_id_to_vendor[emp_id]

        has_ignored_pt = False
        matched_pts = []
        if has_project_type_col and ignore_set:
            pts = emp_rows[FULL_PROJECT_TYPE_COL].dropna().astype(str).str.strip()
            matched = sorted({v for v in pts if v.strip().lower() in ignore_set})
            if matched:
                has_ignored_pt = True
                matched_pts = matched
                project_flagged.append(
                    {
                        "Vendor": vendor,
                        "Emp ID": emp_id,
                        "Employee Name": emp_name,
                        "Total Hours": total_hours,
                        "IgnoredProjectTypes": matched_pts,
                    }
                )
                continue

        employees.append(
            {
                "Vendor": vendor,
                "Emp ID": emp_id,
                "Employee Name": emp_name,
                "Total Hours": total_hours,
                "df": emp_rows,
                "IsUnassigned": is_unassigned,
                "HasIgnoredProjectType": has_ignored_pt,
                "IgnoredProjectTypes": matched_pts,
            }
        )

    return employees, ignored, unassigned, project_flagged, unique_emp_ids


def build_summary_structures(employees, ignored, failed, unassigned, project_flagged, full_df, lang: str):
    now = datetime.now()
    if FULL_DATE_COL in full_df.columns:
        dates = full_df[FULL_DATE_COL].dropna()
        if not dates.empty:
            period_str = f"{dates.min().date()} → {dates.max().date()}"
        else:
            period_str = t("no_dates", lang)
    else:
        period_str = t("no_dates", lang)

    total_emps = len({str(x).strip() for x in full_df[FULL_EMP_ID_COL].dropna().unique()})
    exported_emps = len({e["Emp ID"] for e in employees})
    ignored_emps = len(ignored)
    failed_emps = len(failed)
    unassigned_emps = len({u["Emp ID"] for u in unassigned})
    project_flagged_emps = len({p["Emp ID"] for p in project_flagged})

    summary_stats = {
        "run_timestamp": now,
        "period": period_str,
        "total_emps": total_emps,
        "exported_emps": exported_emps,
        "ignored_emps": ignored_emps,
        "failed_emps": failed_emps,
        "unassigned_emps": unassigned_emps,
        "project_flagged_emps": project_flagged_emps,
    }

    vendor_summary_rows = []
    vendor_group = {}
    for e in employees:
        vendor = e["Vendor"]
        vendor_group.setdefault(vendor, {"hours": 0.0, "count": 0})
        vendor_group[vendor]["hours"] += e["Total Hours"]
        vendor_group[vendor]["count"] += 1

    for vendor, agg in vendor_group.items():
        vendor_summary_rows.append(
            {
                t("table_vendor", lang): vendor,
                t("table_vendor_hours", lang): round(agg["hours"], 2),
                t("table_vendor_emp_count", lang): agg["count"],
            }
        )
    vendor_summary_df = pd.DataFrame(vendor_summary_rows)

    exported_rows = []
    for e in employees:
        exported_rows.append(
            {
                t("table_vendor", lang): e["Vendor"],
                t("table_emp_id", lang): e["Emp ID"],
                t("table_emp_name", lang): e["Employee Name"],
                t("table_emp_hours", lang): round(e["Total Hours"], 2),
            }
        )
    exported_df = pd.DataFrame(exported_rows)

    ignored_rows = []
    for ig in ignored:
        ignored_rows.append(
            {
                t("table_emp_id", lang): ig["Emp ID"],
                t("table_emp_name", lang): ig["Employee Name"],
                t("table_emp_hours", lang): round(ig["Total Hours"], 2),
                t("table_reason", lang): t("reason_not_in_db", lang),
            }
        )
    ignored_df = pd.DataFrame(ignored_rows)

    failed_rows = []
    for fl in failed:
        failed_rows.append(
            {
                t("table_emp_id", lang): fl["Emp ID"],
                t("table_emp_name", lang): fl["Employee Name"],
                t("table_emp_hours", lang): round(fl.get("Total Hours", 0.0), 2),
                t("table_reason", lang): f"{t('reason_exception', lang)}: {fl['Error']}",
            }
        )
    failed_df = pd.DataFrame(failed_rows)

    unassigned_rows = []
    for u in unassigned:
        unassigned_rows.append(
            {
                t("table_vendor", lang): u["Vendor"],
                t("table_emp_id", lang): u["Emp ID"],
                t("table_emp_name", lang): u["Employee Name"],
                t("table_emp_hours", lang): round(u["Total Hours"], 2),
                t("table_unassigned_reason", lang): t("reason_unassigned", lang),
            }
        )
    unassigned_df = pd.DataFrame(unassigned_rows)

    proj_rows = []
    for p in project_flagged:
        proj_rows.append(
            {
                t("table_vendor", lang): p["Vendor"],
                t("table_emp_id", lang): p["Emp ID"],
                t("table_emp_name", lang): p["Employee Name"],
                t("table_emp_hours", lang): round(p["Total Hours"], 2),
                t("table_project_types", lang): ", ".join(p["IgnoredProjectTypes"]),
            }
        )
    project_flagged_df = pd.DataFrame(proj_rows)

    return summary_stats, vendor_summary_df, exported_df, ignored_df, failed_df, unassigned_df, project_flagged_df


def build_docx_summary(
    summary_stats,
    vendor_summary_df,
    exported_df,
    ignored_df,
    failed_df,
    unassigned_df,
    project_flagged_df,
    lang: str,
) -> Document:
    doc = Document()

    doc.add_heading(t("summary_doc_title", lang), level=0)

    doc.add_heading(t("doc_section_overview", lang), level=1)
    ts_str = summary_stats["run_timestamp"].strftime("%Y-%m-%d %H:%M:%S")
    period_str = summary_stats["period"]
    total = summary_stats["total_emps"]
    exported = summary_stats["exported_emps"]
    ignored = summary_stats["ignored_emps"]
    failed = summary_stats["failed_emps"]
    unassigned = summary_stats["unassigned_emps"]
    project_flagged = summary_stats["project_flagged_emps"]

    bullets = [
        t("doc_overview_bullet_1", lang).format(ts=ts_str),
        t("doc_overview_bullet_2", lang).format(period=period_str),
        t("doc_overview_bullet_3", lang).format(total=total),
        t("doc_overview_bullet_4", lang).format(exported=exported),
        t("doc_overview_bullet_5", lang).format(ignored=ignored),
        t("doc_overview_bullet_6", lang).format(failed=failed),
        t("doc_overview_bullet_7", lang).format(unassigned=unassigned),
        t("doc_overview_bullet_8", lang).format(project_flagged=project_flagged),
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    def add_table_section(title_key, df):
        doc.add_heading(t(title_key, lang), level=1)
        if df is not None and not df.empty:
            cols = list(df.columns)
            table = doc.add_table(rows=1 + len(df), cols=len(cols))
            hdr_cells = table.rows[0].cells
            for j, c in enumerate(cols):
                hdr_cells[j].text = str(c)
            for i, (_, row) in enumerate(df.iterrows(), start=1):
                row_cells = table.rows[i].cells
                for j, c in enumerate(cols):
                    row_cells[j].text = str(row[c])
        else:
            doc.add_paragraph("—")

    add_table_section("doc_section_vendor_summary", vendor_summary_df)
    add_table_section("doc_section_exported_emps", exported_df)
    add_table_section("doc_section_ignored", ignored_df)
    add_table_section("doc_section_failed", failed_df)
    add_table_section("doc_section_unassigned", unassigned_df)
    add_table_section("doc_section_project_ignored", project_flagged_df)

    return doc


def build_vendor_staff_summary_df(vendor: str, employees_for_summary: list, lang: str) -> pd.DataFrame:
    col_emp_id = t("table_emp_id", lang)
    col_emp_name = t("table_emp_name", lang)
    col_emp_hours = t("table_emp_hours", lang)

    label_emp_count = t("vendor_staff_header_emp_count", lang)
    label_total_hours = t("vendor_staff_header_total_hours", lang)
    label_avg_hours = t("vendor_staff_header_avg_hours", lang)

    rows = []
    total_hours = 0.0

    for emp in employees_for_summary:
        rows.append(
            {
                col_emp_id: emp["Emp ID"],
                col_emp_name: emp["Employee Name"],
                col_emp_hours: round(emp["Total Hours"], 2),
            }
        )
        total_hours += float(emp["Total Hours"])

    count = len(rows)
    avg_hours = total_hours / count if count else 0.0

    if rows:
        rows.append({col_emp_id: "", col_emp_name: "", col_emp_hours: ""})
        rows.append({col_emp_id: "", col_emp_name: label_emp_count, col_emp_hours: count})
        rows.append({col_emp_id: "", col_emp_name: label_total_hours, col_emp_hours: round(total_hours, 2)})
        rows.append({col_emp_id: "", col_emp_name: label_avg_hours, col_emp_hours: round(avg_hours, 2)})
    else:
        rows.append({col_emp_id: "(no staff)", col_emp_name: "", col_emp_hours: ""})

    return pd.DataFrame(rows)


# ==== Streamlit pages ====

def run_main_page(config: dict, lang: str):
    with st.sidebar:
        st.markdown(f"### ⚙️ {t('settings', lang)}")
        default_mode = config.get("default_output_mode", DEFAULT_CONFIG["default_output_mode"])
        mode_index = 0 if default_mode == "folder" else 1
        output_mode = st.radio(
            t("output_mode", lang),
            ("folder", "zip"),
            index=mode_index,
            format_func=lambda x: t("output_mode_folder", lang) if x == "folder" else t("output_mode_zip", lang),
        )

        default_folder = config.get("default_output_folder", DEFAULT_CONFIG["default_output_folder"])
        output_folder = None
        confirm_clear = False
        if output_mode == "folder":
            output_folder = st.text_input(t("output_folder", lang), value=default_folder)
            if output_folder:
                if os.path.exists(output_folder) and os.listdir(output_folder):
                    st.warning(t("folder_not_empty", lang).format(path=output_folder))
                    confirm_clear = st.checkbox(t("confirm_clear", lang))

        ignore_list_cfg = config.get("ignore_project_types", DEFAULT_IGNORE_LIST)
        st.markdown(f"### 🧾 {t('ignore_list', lang)}")
        st.write(", ".join(ignore_list_cfg) if ignore_list_cfg else "—")

    if not REPORTLAB_AVAILABLE:
        st.warning(t("pdf_not_available", lang))

    db_df = None
    db_loaded_from_config = False
    db_error_msg = None
    db_config_path = config.get("database_path", DEFAULT_CONFIG["database_path"])
    resolved_db_path = resolve_path_from_config(db_config_path) if db_config_path else ""

    if resolved_db_path:
        if os.path.exists(resolved_db_path):
            try:
                db_df = pd.read_excel(resolved_db_path)
                db_loaded_from_config = True
            except Exception as e:
                db_error_msg = f"{t('db_config_error', lang)} {resolved_db_path}\n{e}"
        else:
            db_error_msg = f"{t('db_config_missing', lang)} {resolved_db_path}"

    col1, col2 = st.columns(2)
    with col1:
        full_file = st.file_uploader(t("full_timesheet", lang), type=["xlsx"])
    with col2:
        if db_loaded_from_config:
            st.success(f"{t('db_loaded_from_config', lang)}\n{resolved_db_path}")
            db_file = st.file_uploader(t("db_upload_optional", lang), type=["xlsx"])
        else:
            if db_error_msg:
                st.warning(db_error_msg)
            db_file = st.file_uploader(t("vendor_db", lang), type=["xlsx"])

    if db_file is not None:
        try:
            db_df = pd.read_excel(db_file)
            db_loaded_from_config = False
            db_error_msg = None
        except Exception as e:
            st.error(f"Error reading uploaded Vendor DB: {e}")
            db_df = None

    disable_start = full_file is None or (db_df is None)
    if output_mode == "folder":
        if not output_folder:
            disable_start = True
        elif os.path.exists(output_folder) and os.listdir(output_folder) and not confirm_clear:
            disable_start = True

    start = st.button(t("start", lang), type="primary", disabled=disable_start)

    if not start:
        return

    try:
        try:
            progress = st.progress(0.0, text=t("progress_init", lang))
        except TypeError:
            progress = st.progress(0.0)
            st.info(t("progress_init", lang))

        status_placeholder = st.empty()

        full_df = pd.read_excel(full_file)
        full_df.columns = full_df.columns.str.strip()
        db_df.columns = db_df.columns.str.strip()

        ignore_list = config.get("ignore_project_types", DEFAULT_IGNORE_LIST)

        employees, ignored, unassigned, project_flagged, unique_emp_ids = prepare_employee_data(
            full_df, db_df, ignore_list
        )

        total_to_process = len(employees)
        failed = []
        successful_employees = []

        if output_mode == "folder":
            try:
                if os.path.exists(output_folder):
                    if os.listdir(output_folder):
                        shutil.rmtree(output_folder)
                os.makedirs(output_folder, exist_ok=True)
            except Exception as e:
                st.error(
                    f"Could not clear or create output folder '{output_folder}'. "
                    "Please close any open files (Excel/Explorer) that use this folder and check permissions. "
                    f"Details: {e}"
                )
                return

        zip_buffer = io.BytesIO() if output_mode == "zip" else None
        if zip_buffer is not None:
            zip_file = zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED)
        else:
            zip_file = None

        for idx, emp in enumerate(employees, start=1):
            emp_id = emp["Emp ID"]
            emp_name = emp["Employee Name"]
            vendor = emp["Vendor"]
            emp_df = emp["df"]
            total_hours = emp["Total Hours"]

            try:
                progress_val = (idx / total_to_process) if total_to_process else 1.0
                progress.progress(progress_val)
                status_placeholder.write(
                    t("progress_emp", lang).format(idx=idx, total=total_to_process, emp_id=emp_id)
                )

                safe_vendor_folder = safe_name(vendor)
                file_base_name = f"{vendor}-{emp_id}-{emp_name}.xlsx"
                safe_file_name = safe_name(file_base_name)

                if output_mode == "folder":
                    vendor_folder_path = os.path.join(output_folder, safe_vendor_folder)
                    os.makedirs(vendor_folder_path, exist_ok=True)
                    file_path = os.path.join(vendor_folder_path, safe_file_name)
                    emp_df.to_excel(file_path, index=False)

                    if REPORTLAB_AVAILABLE:
                        pdf_bytes = dataframe_to_pdf_bytes(emp_df, title=file_base_name)
                        pdf_path = (
                            file_path[:-5] + ".pdf"
                            if file_path.lower().endswith(".xlsx")
                            else file_path + ".pdf"
                        )
                        with open(pdf_path, "wb") as pf:
                            pf.write(pdf_bytes)

                else:
                    xls_buffer = io.BytesIO()
                    emp_df.to_excel(xls_buffer, index=False)
                    xls_buffer.seek(0)
                    arcname = f"{safe_vendor_folder}/{safe_file_name}"
                    zip_file.writestr(arcname, xls_buffer.getvalue())

                    if REPORTLAB_AVAILABLE:
                        pdf_bytes = dataframe_to_pdf_bytes(emp_df, title=file_base_name)
                        if safe_file_name.lower().endswith(".xlsx"):
                            pdf_name = safe_file_name[:-5] + ".pdf"
                        else:
                            pdf_name = safe_file_name + ".pdf"
                        pdf_arcname = f"{safe_vendor_folder}/{pdf_name}"
                        zip_file.writestr(pdf_arcname, pdf_bytes)

                successful_employees.append(
                    {
                        "Vendor": vendor,
                        "Emp ID": emp_id,
                        "Employee Name": emp_name,
                        "Total Hours": total_hours,
                    }
                )

            except Exception as e:
                failed.append(
                    {
                        "Vendor": vendor,
                        "Emp ID": emp_id,
                        "Employee Name": emp_name,
                        "Total Hours": total_hours,
                        "Error": str(e),
                    }
                )

        if zip_file is not None:
            zip_file.close()

        employees_for_summary = []
        success_keys = {(e["Emp ID"], e["Vendor"]) for e in successful_employees}
        for emp in employees:
            if (emp["Emp ID"], emp["Vendor"]) in success_keys:
                employees_for_summary.append(emp)

        unassigned_for_summary = [
            u for u in unassigned if (u["Emp ID"], u["Vendor"]) in success_keys
        ]

        project_flagged_for_summary = project_flagged

        (
            summary_stats,
            vendor_summary_df,
            exported_df,
            ignored_df,
            failed_df,
            unassigned_df,
            project_flagged_df,
        ) = build_summary_structures(
            employees_for_summary,
            ignored,
            failed,
            unassigned_for_summary,
            project_flagged_for_summary,
            full_df,
            lang,
        )

        vendor_to_emps = {}
        for emp in employees_for_summary:
            v = emp["Vendor"]
            vendor_to_emps.setdefault(v, []).append(emp)

        if output_mode == "folder":
            for vendor, emps_list in vendor_to_emps.items():
                safe_vendor_folder = safe_name(vendor)
                vendor_folder_path = os.path.join(output_folder, safe_vendor_folder)
                os.makedirs(vendor_folder_path, exist_ok=True)

                vendor_summary_df_vendor = build_vendor_staff_summary_df(vendor, emps_list, lang)
                base_name = safe_name(f"{vendor}-StaffSummary")

                vendor_summary_xlsx_path = os.path.join(vendor_folder_path, base_name + ".xlsx")
                vendor_summary_df_vendor.to_excel(vendor_summary_xlsx_path, index=False)

                if REPORTLAB_AVAILABLE:
                    pdf_bytes = dataframe_to_pdf_bytes(vendor_summary_df_vendor, title=f"{vendor} - Staff Summary")
                    vendor_summary_pdf_path = os.path.join(vendor_folder_path, base_name + ".pdf")
                    with open(vendor_summary_pdf_path, "wb") as pf:
                        pf.write(pdf_bytes)

        else:
            if zip_buffer is not None:
                with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file_append:
                    for vendor, emps_list in vendor_to_emps.items():
                        safe_vendor_folder = safe_name(vendor)
                        vendor_summary_df_vendor = build_vendor_staff_summary_df(vendor, emps_list, lang)
                        base_name = safe_name(f"{vendor}-StaffSummary")

                        xls_buf = io.BytesIO()
                        vendor_summary_df_vendor.to_excel(xls_buf, index=False)
                        xls_buf.seek(0)
                        arcname_xlsx = f"{safe_vendor_folder}/{base_name}.xlsx"
                        zip_file_append.writestr(arcname_xlsx, xls_buf.getvalue())

                        if REPORTLAB_AVAILABLE:
                            pdf_bytes = dataframe_to_pdf_bytes(
                                vendor_summary_df_vendor, title=f"{vendor} - Staff Summary"
                            )
                            arcname_pdf = f"{safe_vendor_folder}/{base_name}.pdf"
                            zip_file_append.writestr(arcname_pdf, pdf_bytes)

        doc = build_docx_summary(
            summary_stats,
            vendor_summary_df,
            exported_df,
            ignored_df,
            failed_df,
            unassigned_df,
            project_flagged_df,
            lang,
        )
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        if output_mode == "folder":
            summary_path = os.path.join(output_folder, SUMMARY_DOC_NAME)
            with open(summary_path, "wb") as f:
                f.write(doc_buffer.getvalue())
        else:
            zip_buffer.seek(0)
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zip_file_append:
                zip_file_append.writestr(SUMMARY_DOC_NAME, doc_buffer.getvalue())
            zip_buffer.seek(0)

        progress.progress(1.0)
        status_placeholder.empty()
        st.success(t("done", lang))

        st.markdown(f"### 📊 {t('metrics_title', lang)}")
        c1, c2, c3, c4, c5, c6 = st.columns(6)
        c1.metric(t("metric_total_emps", lang), summary_stats["total_emps"])
        c2.metric(t("metric_exported_emps", lang), summary_stats["exported_emps"])
        c3.metric(t("metric_ignored_emps", lang), summary_stats["ignored_emps"])
        c4.metric(t("metric_failed_emps", lang), summary_stats["failed_emps"])
        c5.metric(t("metric_unassigned_emps", lang), summary_stats["unassigned_emps"])
        c6.metric(
            t("metric_project_flagged_emps", lang),
            summary_stats["project_flagged_emps"],
        )

        st.markdown("----")
        c7, c8 = st.columns(2)
        c7.write(
            f"**{t('run_timestamp', lang)}:** {summary_stats['run_timestamp'].strftime('%Y-%m-%d %H:%M:%S')}"
        )
        c8.write(f"**{t('period', lang)}:** {summary_stats['period']}")

        with st.expander("📦 " + t("vendor_summary_title", lang), expanded=True):
            if vendor_summary_df is not None and not vendor_summary_df.empty:
                st.dataframe(vendor_summary_df, use_container_width=True)
            else:
                st.write("—")

        with st.expander("👤 " + t("exported_table_title", lang), expanded=False):
            if exported_df is not None and not exported_df.empty:
                st.dataframe(exported_df, use_container_width=True)
            else:
                st.write("—")

        with st.expander("👀 " + t("ignored_table_title", lang), expanded=False):
            if ignored_df is not None and not ignored_df.empty:
                st.dataframe(ignored_df, use_container_width=True)
            else:
                st.write(t("no_ignored", lang))

        with st.expander("🧩 " + t("unassigned_table_title", lang), expanded=False):
            if unassigned_df is not None and not unassigned_df.empty:
                st.dataframe(unassigned_df, use_container_width=True)
            else:
                st.write(t("no_unassigned", lang))

        with st.expander("🚩 " + t("project_ignored_table_title", lang), expanded=False):
            if project_flagged_df is not None and not project_flagged_df.empty:
                st.dataframe(project_flagged_df, use_container_width=True)
            else:
                st.write(t("no_project_flagged", lang))

        with st.expander("⚠️ " + t("failed_table_title", lang), expanded=False):
            if failed_df is not None and not failed_df.empty:
                st.dataframe(failed_df, use_container_width=True)
            else:
                st.write(t("no_failed", lang))

        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                label="📄 " + t("download_summary", lang),
                data=doc_buffer,
                file_name=SUMMARY_DOC_NAME,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_dl2:
            if output_mode == "zip" and zip_buffer is not None:
                st.download_button(
                    label="🗂️ " + t("download_zip", lang),
                    data=zip_buffer,
                    file_name=f"timesheet_split_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                    mime="application/zip",
                )

    except Exception as e:
        st.error(f"{t('fatal_error', lang)}: {e}")


def run_settings_page(config: dict, lang: str):
    st.markdown(f"## {t('config_title', lang)}")

    if st.button("⬅️ Back"):
        st.session_state["page"] = "main"
        return

    cfg = config.copy()

    st.markdown(f"### 📁 {t('config_section_paths', lang)}")
    db_path = st.text_input(
        t("config_db_path", lang),
        value=cfg.get("database_path", DEFAULT_CONFIG["database_path"]),
    )
    logo_path = st.text_input(
        t("config_logo_path", lang),
        value=cfg.get("logo_path", DEFAULT_CONFIG["logo_path"]),
    )

    if db_path:
        resolved = resolve_path_from_config(db_path)
        st.caption(f"{t('config_db_resolved', lang)}: {resolved}")
    else:
        resolved = ""

    st.markdown(f"### 🎨 {t('config_section_branding', lang)}")
    dept_name = st.text_input(
        t("config_department_name", lang),
        value=cfg.get("department_name", DEFAULT_CONFIG["department_name"]),
    )
    user_name = st.text_input(
        t("config_user_name", lang),
        value=cfg.get("user_name", DEFAULT_CONFIG["user_name"]),
    )

    st.markdown(f"**{t('config_logo_preview', lang)}:**")
    logo_resolved = resolve_path_from_config(logo_path) if logo_path else ""
    if logo_resolved and os.path.exists(logo_resolved):
        st.image(logo_resolved, width=150)
    else:
        st.caption("No logo found at current path.")

    st.markdown(f"### 📤 {t('config_section_output', lang)}")
    default_mode = cfg.get("default_output_mode", DEFAULT_CONFIG["default_output_mode"])
    mode_index = 0 if default_mode == "folder" else 1
    out_mode = st.radio(
        t("config_default_output_mode", lang),
        ("folder", "zip"),
        index=mode_index,
        format_func=lambda x: t("output_mode_folder", lang) if x == "folder" else t("output_mode_zip", lang),
    )
    out_folder = st.text_input(
        t("config_default_output_folder", lang),
        value=cfg.get("default_output_folder", DEFAULT_CONFIG["default_output_folder"]),
    )

    st.markdown(f"### 🚫 {t('config_section_ignore', lang)}")
    ignore_list = cfg.get("ignore_project_types", DEFAULT_IGNORE_LIST)
    ignore_df = pd.DataFrame({"Project Type": ignore_list})
    edited_ignore_df = st.data_editor(
        ignore_df,
        num_rows="dynamic",
        key="ignore_editor",
    )

    if st.button(t("config_save_button", lang)):
        new_cfg = cfg.copy()
        new_cfg["database_path"] = db_path.strip() or DEFAULT_CONFIG["database_path"]
        new_cfg["logo_path"] = logo_path.strip() or DEFAULT_CONFIG["logo_path"]
        new_cfg["department_name"] = dept_name.strip() or DEFAULT_CONFIG["department_name"]
        new_cfg["user_name"] = user_name.strip() or DEFAULT_CONFIG["user_name"]
        new_cfg["default_output_mode"] = out_mode
        new_cfg["default_output_folder"] = out_folder.strip() or DEFAULT_CONFIG["default_output_folder"]

        try:
            new_ignore_list = [
                str(v).strip()
                for v in edited_ignore_df["Project Type"].tolist()
                if str(v).strip()
            ]
        except Exception:
            new_ignore_list = DEFAULT_IGNORE_LIST
        if not new_ignore_list:
            new_ignore_list = DEFAULT_IGNORE_LIST
        new_cfg["ignore_project_types"] = new_ignore_list

        save_app_config(new_cfg)
        st.success(t("config_saved", lang))


def main():
    st.set_page_config(page_title="Timesheet Splitter", page_icon="⏱️", layout="wide")
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

    config = load_app_config()

    # Initialize language from config
    if "lang" not in st.session_state:
        cfg_lang = config.get("language", DEFAULT_CONFIG["language"])
        if isinstance(cfg_lang, str) and "ar" in cfg_lang.lower():
            st.session_state["lang"] = "ar"
        else:
            st.session_state["lang"] = "en"

    if "page" not in st.session_state:
        st.session_state["page"] = "main"

    with st.sidebar:
        lang_choice = st.radio(
            t("sidebar_lang", st.session_state["lang"]),
            ("en", "ar"),
            index=0 if st.session_state["lang"] == "en" else 1,
            format_func=lambda x: t(f"ui_language_{x}", st.session_state["lang"]),
        )
        if lang_choice != st.session_state["lang"]:
            st.session_state["lang"] = lang_choice
            config["language"] = "Arabic" if lang_choice == "ar" else "English"
            save_app_config(config)

    lang = st.session_state["lang"]

    logo_path_cfg = config.get("logo_path", DEFAULT_CONFIG["logo_path"])
    logo_full_path = resolve_path_from_config(logo_path_cfg) if logo_path_cfg else ""
    dept = config.get("department_name", DEFAULT_CONFIG["department_name"])
    user_name = config.get("user_name", DEFAULT_CONFIG["user_name"])

    col_logo, col_title, col_gear = st.columns([1, 4, 0.7])
    with col_logo:
        if logo_full_path and os.path.exists(logo_full_path):
            st.image(logo_full_path, use_container_width=True)
        else:
            st.write("")
    with col_title:
        st.markdown(f"<h1 class='app-title'>⏱️ {t('title', lang)}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p class='app-subtitle'>{t('subtitle', lang)}</p>", unsafe_allow_html=True)
        st.markdown(
            f"<div class='brand-line'><strong>Department:</strong> {dept} &nbsp;&nbsp;|&nbsp;&nbsp; "
            f"<strong>User:</strong> {user_name}</div>",
            unsafe_allow_html=True,
        )
        if st.session_state["page"] == "main":
            st.markdown(
                f"<div class='welcome-line'>👋 Welcome, {user_name}</div>",
                unsafe_allow_html=True,
            )
    with col_gear:
        if st.button("⚙️", help="Open settings"):
            st.session_state["page"] = "settings"

    if st.session_state["page"] == "main":
        run_main_page(config, lang)
    else:
        run_settings_page(config, lang)


if __name__ == "__main__":
    main()
